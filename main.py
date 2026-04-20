import re
import uuid

import enchant
import numpy as np
import unicodedata
from datetime import datetime, timezone
from dataclasses import dataclass
from typing import Dict, List, Sequence, Optional, Tuple
from lingua import Language, LanguageDetectorBuilder
from sympy.codegen.ast import none
from googletrans import Translator
from collections import defaultdict
from collections import Counter

import fitz
import pdfplumber
import fasttext
from docx import Document


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
ZH_PREFIXES = ("zh", "ja", "ko")
IGNORE_TERMS = {"unk", "n/a", "na", "ni", "yes", "no", "fr", "comment", "comments", "sender",
                "test", "result", "value", "qualifier", "date", "name", "type", "code",
                "status", "number", "source", "report", "case", "drug", "dose", "route", "senders",
                "given", "value/qualifier", "duplicate", "textual", "medra", "unit", "paclitaxel arrow",
                "exfumador", "transminases", "afssaps", "fr-afssaps-", "france",
                }
COMMON_MEDICAL_TERMS = {
    "lipase", "glucose", "albumin", "protein", "plasma",
    "insulin", "cholesterol", "bilirubin", "carcinoma", "paxlovid", "pembrolizumab", "patient",
}
# Common English words that look French/Italian to fastText
# (pharma regulatory domain cognates + British spellings)
_ENGLISH_COGNATES = {
    "authorisation", "application", "indication", "transmission",
    "identification", "information", "relevant", "dosage", "version",
    "country", "number", "batch", "wrapper", "and", "for", "of",
    "the", "with", "meddra", "icsr", "ich", " hospitalisation", "available", "congenital",
    "anomaly", "birth", "defect", "organisation", "therapies"
}

def _all_words_english(text: str) -> bool:
    """True if every alphabetic token is a known English/cognate word."""
    words = re.findall(r"[^\W\d_]+", text.lower(), re.UNICODE)
    if not words:
        return False
    return all(w in _ENGLISH_COGNATES for w in words)

# ---------------------------------------------------------------------------
# Dataclasses  (FIX 5: MatchHit and WordBox were used but never defined)
# ---------------------------------------------------------------------------

@dataclass
class BboxPdf:
    x0: float
    y0: float
    x1: float
    y1: float

    def as_tuple(self) -> Tuple[float, float, float, float]:
        return (self.x0, self.y0, self.x1, self.y1)


@dataclass
class WordBox:
    token: str
    bbox_pdf: BboxPdf
    page_index_1based: int = 1


# ---------------------------------------------------------------------------
# Model loader  (FIX 2: lazy load — don't crash at import if file missing)
# ---------------------------------------------------------------------------

_language_model = None

def get_language_model(model_path: str = "lid.176.bin"):
    global _language_model
    if _language_model is None:
        _language_model = fasttext.load_model(model_path)
    return _language_model


# ---------------------------------------------------------------------------
# Document metadata
# ---------------------------------------------------------------------------

def build_document_metadata(file_name, file_type, source_lang=None, foreign_lang=None):
    return {
        "document_id": str(uuid.uuid4()),
        "file_name": file_name,
        "file_type": file_type,
        "from_language": foreign_lang,
        "to_language": source_lang,
        "process_stage": "uploaded",
        "created_at": datetime.now(timezone.utc).isoformat()
    }


# ---------------------------------------------------------------------------
# Parsers
# ---------------------------------------------------------------------------

def parse_txt(file_path: str) -> dict:
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()
    return {
        "text": text,
        "lines": text.splitlines()
    }


def parse_docx(file_path: str) -> dict:
    doc = Document(file_path)

    # Collect table cells
    table_data = []
    for table_no, table in enumerate(doc.tables, start=1):
        for row_no, row in enumerate(table.rows, start=1):
            for column_no, cell in enumerate(row.cells, start=1):
                table_data.append({
                    "table_no":  table_no,
                    "row_no":    row_no,
                    "column_no": column_no,
                    "cell_text": cell.text or ""
                })

    # Track paragraphs that belong to table cells so we don't double-count them
    table_paragraphs = set()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    table_paragraphs.add(id(para))

    # Collect only standalone paragraphs (not inside tables), skip empty ones
    paragraphs = []
    for idx, para in enumerate(doc.paragraphs, start=1):
        if id(para) in table_paragraphs:
            continue
        if para.text.strip():
            paragraphs.append({
                "paragraph_no": idx,
                "text":         para.text
            })

    return {
        "paragraph_data": paragraphs,
        "table_data":     table_data
    }


def parse_pdf(file_path: str) -> dict:
    doc = fitz.open(file_path)

    # Collect table cells via pdfplumber
    extracted_tables = []
    table_bboxes = []   # track table regions so we can exclude them from text spans

    with pdfplumber.open(file_path) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            tables = page.find_tables()
            if not tables:
                continue
            for table_no, table in enumerate(tables, start=1):
                # Store table bbox to exclude from text extraction later
                table_bboxes.append((page_number, table.bbox))
                for row_no, row in enumerate(table.rows, start=1):
                    for column_no, cell in enumerate(row.cells, start=1):
                        if cell is None:
                            continue
                        x0, top, x1, bottom = cell
                        if any(v is None for v in (x0, top, x1, bottom)):
                            continue

                        cell_text = page.crop((x0, top, x1, bottom)).extract_text() or ""
                        # if any(w in cell_text for w in ["Femenina", "Vinculada", "marzo", "Exfumador", "Cáncer"]):
                        #     print(f"Page {page_number} | Table {table_no} | Row {row_no} | Col {column_no}")
                        #     print(f"  bbox: ({x0:.1f}, {top:.1f}, {x1:.1f}, {bottom:.1f})")
                        #     print(f"  text: '{cell_text}'")
                        extracted_tables.append({
                            "page":      page_number,
                            "table_no":  table_no,
                            "row_no":    row_no,
                            "column_no": column_no,
                            "cell_text": cell_text,
                            "x0":        x0,
                            "y0":        top,
                            "x1":        x1,
                            "y1":        bottom
                        })

    # Build a set of (page, bbox) for quick lookup
    def span_in_table(page_no, bbox) -> bool:
        sx0, sy0, sx1, sy1 = bbox
        for (tpage, (tx0, ty0, tx1, ty1)) in table_bboxes:
            if tpage != page_no:
                continue
            # Check if span overlaps with table bbox
            if sx0 >= tx0 - 1 and sy0 >= ty0 - 1 and sx1 <= tx1 + 1 and sy1 <= ty1 + 1:
                return True
        return False

    # Collect text spans that are NOT inside a table region
    extracted_text = []
    for page_no, page in enumerate(doc, start=1):
        blocks = page.get_text("dict")["blocks"]
        for block_idx, block in enumerate(blocks, start=1):
            for line_idx, line in enumerate(block.get("lines", []), start=1):
                for span in line.get("spans", []):
                    if span_in_table(page_no, span["bbox"]):
                        continue   # already captured in table_data
                    extracted_text.append({
                        "page":     page_no,
                        "block_no": block_idx,
                        "line_no":  line_idx,
                        "text":     span["text"],
                        "bbox":     span["bbox"]
                    })

    return {
        "text_data":  extracted_text,
        "table_data": extracted_tables
    }


# FIX 1: removed async — none of the parsers are async, no await needed
def process_file(input_file: str, src_lang: str = "hi", target_lang: str = "en") -> dict:
    if input_file.endswith(".txt"):
        return parse_txt(input_file)
    elif input_file.endswith(".docx"):
        return parse_docx(input_file)
    elif input_file.endswith(".pdf"):
        return parse_pdf(input_file)
    else:
        raise ValueError(f"Unsupported file format: {input_file}")


# ---------------------------------------------------------------------------
# Normalization
# ---------------------------------------------------------------------------

def normalize(text: str, lowercase: bool = False) -> str:
    # Unicode normalize — converts é→e, ê→e, etc.
    text = unicodedata.normalize("NFKD", text)
    text = "".join(c for c in text if not unicodedata.combining(c))
    # Collapse whitespace
    text = re.sub(r"\s+", " ", text).strip()
    # FIX 9: do NOT strip / or . — needed for MG/J, S. aureus, CHO cells
    # Only strip characters that are never part of a pharmaceutical term
    text = re.sub(r"[\"\'\\|<>]", "", text)
    if lowercase:
        text = text.lower()
    return text

def normalize_for_detection(text: str) -> str:
    text = re.sub(r"\s+", " ", text).strip()
    text = re.sub(r"[\"\'\\|<>]", "", text)
    return text


# ---------------------------------------------------------------------------
# Pre-detection filters
# ---------------------------------------------------------------------------

# Matches case ID codes like FR-AFSSAPS-BX2024002066, IT-MINISAL02-1009857
_CASE_ID_RE = re.compile(r'^[A-Z]{2,3}-[A-Z0-9]+-[A-Z0-9]+$', re.IGNORECASE)

# Matches pure acronyms — all caps, no spaces, 2-15 chars
_ACRONYM_RE = re.compile(r'^[A-Z]{2,15}$')

# Matches organisation names like "AIFA - Agenzia Italiana del Farmaco"
_ORG_NAME_RE = re.compile(r'^[A-Z]{2,6}\s*[-–]\s*.+$')

# Matches dosage-only strings like "200 Mg milligram(s)", "5 AUC", "70 Years"
_DOSAGE_RE = re.compile(
    r'^[\d\s\.,]+\s*'
    r'(mg|mcg|ug|ml|l|auc|kg|g|mmol|nmol'
    r'|milligram|microgram|liter|litre'
    r'|tablet|capsule|vial|ampoule'
    r'|dose|unit|units|interval|days|weeks|months|years)'
    r'[\s\(\)a-z\d\.,/]*$',
    re.IGNORECASE
)

_STRUCTURED_ID_RE = re.compile(r'^[A-Z]{2,3}\s*/\s*\d+(?:\s*/\s*\d+)+$', re.IGNORECASE)

# All-caps section headers like "PATIENT DEATH", "RELEVANT PAST DRUG HISTORY"
# Excludes strings with parentheses — those may be substance descriptions
_ALL_CAPS_HEADER_RE = re.compile(r'^[A-Z][A-Z\s\/\-]+$')

# # Field labels ending with ? or : like "Was Autopsy Done?", "Test Result (unit):"
# _FIELD_LABEL_RE = re.compile(r'^.+[\?\:]$')

#
# # Strips a leading drug name + dosage: "KEYTRUDA 25 mg/mL, " → ""
# _DRUG_PREFIX_RE = re.compile(
#     r'(mg|mcg|ml|g|ug|mmol|mg/ml|mcg/ml)'
#     r'[^,]*,?\s*',
#     re.IGNORECASE
# )
_NUMBER_TOKEN_RE = re.compile(
    r'^[\d\.\,%\+\-]+$'         # pure numbers, percentages, decimals: 3, 12%, 40%
    r'|^[\W_]+$'                # only punctuation (..., /, •)
    r'|^/?r[\W_]*$'
    r'|^\d+[a-zA-Z]{1,3}$',     # number+unit suffix: 25mg, 200ml
    re.IGNORECASE
)

def _extract_detectable_segment(text: str) -> str:
    """Prefer content inside parentheses; also replace slashes with spaces."""
    match = re.search(r'\(+([^)]+)\)+', text)
    if match:
        inner = match.group(1).strip()
        if len(inner) > 5:
            return inner.replace("/", " ")
    return text.replace("/", " ")

def _remove_numbers(text: str) -> str:
    # remove standalone numbers (years, counts, etc.)
    text = re.sub(r"\b\d{1,4}\b", " ", text)

    # remove date-like patterns (18/03/2020, 18-03-2020)
    text = re.sub(r"\b\d{1,2}[-/]\d{1,2}[-/]\d{2,4}\b", " ", text)

    # remove numeric + punctuation noise (e.g., 12%, 3.5, etc.)
    text = re.sub(r"\b\d+[\d\.,%]*\b", " ", text)

    return re.sub(r"\s+", " ", text).strip()

def prepare_for_detection(text: str) -> Optional[str]:
    """
    Full pre-processing pipeline before language detection.
    Returns None if the text should be skipped entirely.
    """
    # text = normalize_for_detection(text)

    if not text or text.lower().strip() in IGNORE_TERMS:
        #print(f"  KILLED by IGNORE_TERMS: '{text}'")
        return None
    stripped = text.strip()
    if _CASE_ID_RE.match(stripped):
        #print(f"  KILLED by CASE_ID_RE: '{text}'")
        return None
    if _ACRONYM_RE.match(stripped):
        # print(f"  KILLED by ACRONYM_RE: '{text}'")
        return None
    if _ORG_NAME_RE.match(stripped):
        # print(f"  KILLED by ORG_NAME_RE: '{text}'")
        return None
    if _STRUCTURED_ID_RE.match(stripped):
        print(f"  KILLED by STRUCTURED_ID_RE: '{text}'")
        return None
    if _ALL_CAPS_HEADER_RE.match(stripped) and '(' not in stripped:
        # print(f"  KILLED by ALL_CAPS_HEADER_RE: '{text}'")
        return None
    if _DOSAGE_RE.match(text.strip()):
        # print(f"  KILLED by DOSAGE_RE: '{text}'")
        return None
    # 🔥 ADD HERE
    text = _remove_numbers(text)
    if len(text.strip()) < 2:
        print(f"  KILLED by length: '{text}'")
        return None

    return text.strip().strip(":")


def normalize_for_match(text: str) -> str:
    """Normalize for glossary key comparison — always lowercase."""
    return normalize(text, lowercase=True)


def normalize_word_for_match(word: str) -> str:
    """Normalize a single token for Aho-Corasick matching."""
    return normalize_for_match(word)

# Build once at startup — restrict to languages you actually expect
# This makes it faster and more accurate than detecting all 75 languages
_lingua_detector = None

def get_lingua_detector():
    global _lingua_detector
    if _lingua_detector is None:
        _lingua_detector = (
            LanguageDetectorBuilder
            .from_languages(
                Language.FRENCH,
                #Language.GERMAN,
                #Language.SPANISH,
                Language.ENGLISH,
                #Language.ITALIAN
            )
            .with_minimum_relative_distance(0.11)
            .build()
        )
    return _lingua_detector

# ---------------------------------------------------------------------------
# Language detection
# ---------------------------------------------------------------------------



def detect_language_by_tokens(text: str, model_path: str):
    detector = get_lingua_detector()
    model = get_language_model(model_path)

    scores = defaultdict(float)

    # Unicode-aware tokenization: match "letters" across scripts (Latin, Hangul, Devanagari, etc.)
    # This intentionally excludes digits/underscore and most punctuation.
    tokens = re.findall(r"\b[^\W\d_]+\b", text, flags=re.UNICODE)

    for token in tokens:
        token_clean = token.strip().lower()

        # skip very short tokens
        if len(token_clean) <= 2:
            continue

        # skip uppercase acronyms like CHO, DNA, HER
        # Only apply this heuristic for ASCII words; non-Latin scripts don't have the same casing concept.
        if re.fullmatch(r"[A-Z]{2,4}", token):
            continue

        # ----- Try Lingua first -----
        lang = detector.detect_language_of(token_clean)

        if lang is not None:
            conf = detector.compute_language_confidence(token_clean, lang)

            if conf >= 0.85:
                lang_code = lang.iso_code_639_1.name.lower()
                weight = conf * min(len(token_clean) / 5, 2)
                scores[lang_code] += weight
                continue

        # ----- Fallback to fastText -----
        labels, probs = model.predict(token_clean, k=1)

        ft_lang = labels[0].replace("__label__", "")
        ft_conf = float(probs[0])

        if ft_conf >= 0.65:
            weight = ft_conf * min(len(token_clean) / 5, 2)
            scores[ft_lang] += weight

    if not scores:
        return None

        # ✅ normalization FIX
    best_lang = max(scores, key=scores.get)
    best_score = scores[best_lang]
    total_score = sum(scores.values())

    normalized_conf = best_score / total_score if total_score else 0

    if normalized_conf < 0.65:
        return None

    return {
        "language": best_lang,
        "confidence": normalized_conf
    }

# _xlm_pipeline = None
#
# def get_xlm_pipeline():
#     global _xlm_pipeline
#     if _xlm_pipeline is None:
#         _xlm_pipeline = hf_pipeline(
#             "text-classification",
#             model="papluca/xlm-roberta-base-language-detection",
#             device=-1  # CPU; change to 0 if you have GPU
#         )
#     return _xlm_pipeline
#
#
# def detect_with_xlmroberta(text: str) -> Optional[dict]:
#     try:
#         result = get_xlm_pipeline()(text[:512], truncation=True)[0]
#         return {
#             "language":   result["label"],
#             "confidence": round(result["score"], 4)
#         }
#     except Exception:
#         return None

def detect_language(text: str, model_path: str = r"D:\TMS\Models\lid.176.bin") -> Optional[dict]:
    if not text or not text.strip():
        print("no textt")
        return None


    text_clean = text.replace("\n", " ").strip()
    text_lower  = text_clean.lower()
    detector    = get_lingua_detector()
    words       = text_clean.split()

    # print(f"INPUT: '{text}' | CLEAN: '{text_clean}' | LEN: {len(text_clean)}")

    # --- Single word: xlm-roberta only ---
    # Lingua misclassifies single English words as foreign (Handwritten→nl, Register→ur)
    #xlm-roberta is more reliable here but needs a high confidence bar
    # if len(words) <= 3:
    #     result = detect_with_xlmroberta(text_clean)
    #     # Non-Latin scripts (ko/hi/ja/zh/ru/...) are usually very confident even for short text.
    #     # Keep English false-positives low, but avoid missing short Korean/Hindi phrases.
    #     if result and result["confidence"] >= 0.99:
    #         return result
    #     detector = get_lingua_detector()
    #     lang = detector.detect_language_of(text_lower)
    #
    #     if lang is not None:
    #         conf = detector.compute_language_confidence(text_lower, lang)
    #
    #         if conf >= 0.85:
    #             return {
    #                 "language": lang.iso_code_639_1.name.lower(),
    #                 "confidence": round(conf, 4)
    #             }
    #
    #     return None

    # --- Short multi-word (< 20 chars): Lingua → xlm-roberta ---
    if len(text_clean) < 20:
        lang = detector.detect_language_of(text_lower)
        print("whyyyyyyyy noneeeeeeeee", text_lower, lang)
        if lang is not None:
            conf = detector.compute_language_confidence(text_lower, lang)
            if conf >= 0.50:
                print(text_clean)
                print("hereeeeeee", lang)
                return {"language": lang.iso_code_639_1.name.lower(), "confidence": round(conf, 4)}
        # return detect_with_xlmroberta(text_clean)

        return None

    # --- Longer text: fastText → xlm-roberta → Lingua ---
    model   = get_language_model(model_path)
    labels, probs = model.predict(text_lower, k=3)
    probs   = np.asarray(probs)
    ft_lang = labels[0].replace("__label__", "")
    ft_conf = round(float(probs[0]), 4)

    if ft_conf >= 0.75:
        return {"language": ft_lang, "confidence": ft_conf}

    lang = detector.detect_language_of(text_lower)
    if lang is not None:
        conf = detector.compute_language_confidence(text_lower, lang)
        return {"language": lang.iso_code_639_1.name.lower(), "confidence": round(conf, 4)}

    # Both fastText and Lingua uncertain — use xlm-roberta as last resort
    # Only for multi-word text (single words handled separately above)
    # if len(words) >= 2:
    #     xlm_result = detect_with_xlmroberta(text_clean)
    #     if xlm_result and xlm_result["confidence"] >= 0.99:
    #         return xlm_result

    if 0.20 <= ft_conf <= 0.70:
        # --- FINAL fallback: token-level detection ---
        print("lowwwww conffffff", text_clean)
        token_result = detect_language_by_tokens(text_clean, model_path)

        if token_result:
            return token_result

        return None

    return None

def is_probably_english_word(token: str) -> bool:
    """Rough heuristic — pure ASCII letters are likely English."""
    return bool(re.match(r"^[A-Za-z]+$", token))

# ---------------------------------------------------------------------------
# Foreign content extraction
# ---------------------------------------------------------------------------

def extract_foreign_content(
    parsed_data: dict,
    base_language: str = "en"
) -> List[dict]:

    foreign_records = []
    model_path = r"D:\TMS\Models\lid.176.bin"

    # def is_foreign(lang_result: dict, text: str) -> bool:
    #     if lang_result["language"] == "en":
    #         return False
    #
    #     words = re.findall(r"[A-Za-zÀ-ÿ]+", text)
    #
    #     if len(words) <= 2:
    #         return lang_result["confidence"] >= 0.85
    #
    #     if len(words) <= 5:
    #         return lang_result["confidence"] >= 0.70
    #
    #     return lang_result["confidence"] >= 0.55

    def is_foreign(lang_result: dict, text: str) -> bool:
        if lang_result["language"] == base_language:
            return False
        # # Reject obvious English tokens
        # if is_probably_english_word(text):
        #     print("probably english word", text)
        #     return False
            # If every token is a known English word, don't flag regardless of model confidence
        if _all_words_english(text):
            return False
        min_conf = 0.85 if len(text.strip()) < 20 else 0.78
        result =  lang_result["confidence"] >= min_conf
        if result:
            print(f"  FLAGGED: '{text[:60]}' lang={lang_result['language']} conf={lang_result['confidence']:.2f}")
        return result

    # PDF span-wise
    for item in parsed_data.get("text_data", []):
        text = item.get("text", "").strip()
        detect_text = prepare_for_detection(text)
        print("prepare result:", detect_text)

        if not detect_text:
            continue

        lang_result = detect_language(detect_text, model_path)
        print("lang result:", lang_result)

        #print("is_foreign:", is_foreign(lang_result, detect_text) if lang_result else "N/A")
        # Always run fallback as cross-check, not just on failure
        if not lang_result:
            continue
        # STEP 1 — trust strong primary result
        if is_foreign(lang_result, detect_text):
            pass  # keep it

        else:
            fallback_text = _extract_detectable_segment(detect_text)
            if fallback_text != detect_text:
                fallback_result = detect_language(fallback_text, model_path)
                if fallback_result and is_foreign(fallback_result, fallback_text):
                    lang_result = fallback_result
                    detect_text = fallback_text
                else:
                    token_result = detect_language_by_tokens(detect_text, model_path)
                    if token_result and is_foreign(token_result, detect_text):
                        # token-level agrees there's something foreign
                        lang_result = token_result
                    else:
                        # both fallbacks say not foreign — trust them over primary
                        continue
            else:
                token_result = detect_language_by_tokens(detect_text, model_path)
                if token_result and is_foreign(token_result, detect_text):
                    lang_result = token_result
                else:
                    # token-level says not foreign — discard primary result
                    continue


        # # ADD THIS BLOCK
        # token_result = detect_language_by_tokens(detect_text, model_path)
        #
        # if token_result:
        #     if token_result["language"] != lang_result["language"]:
        #         continue  # disagreement → skip

        bbox = item.get("bbox", [None, None, None, None])
        foreign_records.append({
            "source_type":          "text",
            "word":                 text,
            "detect_word": detect_text,
            "source_language":      lang_result["language"],
            "language_confidence":  lang_result["confidence"],
            "page_no":              item.get("page"),
            "line_no":              item.get("line_no"),
            "x0":                   bbox[0],
            "y0":                   bbox[1],
            "x1":                   bbox[2],
            "y1":                   bbox[3],
        })

    # DOCX paragraph-wise
    for item in parsed_data.get("paragraph_data", []):
        text = item.get("text", "").strip()
        detect_text = prepare_for_detection(text)
        #print("prepare result:", detect_text)
        if not detect_text:
            continue

        lang_result = detect_language(detect_text, model_path)
        # print("lang result:", lang_result)

        # print("is_foreign:", is_foreign(lang_result, detect_text) if lang_result else "N/A")
        if not lang_result:
            continue
        # STEP 1 — trust strong primary result
        if is_foreign(lang_result, detect_text):
            pass  # keep it
        # Always run fallback as cross-check, not just on failure
        else:
            fallback_text = _extract_detectable_segment(detect_text)

            if fallback_text != detect_text:
                fallback_result = detect_language(fallback_text, model_path)
                if fallback_result and is_foreign(fallback_result, fallback_text):
                    lang_result = fallback_result
                    detect_text = fallback_text
                else:
                    token_result = detect_language_by_tokens(detect_text, model_path)
                    if token_result and is_foreign(token_result, detect_text):
                        # token-level agrees there's something foreign
                        lang_result = token_result
                    else:
                        # both fallbacks say not foreign — trust them over primary
                        continue
            else:
                token_result = detect_language_by_tokens(detect_text, model_path)
                if token_result and is_foreign(token_result, detect_text):
                    lang_result = token_result
                else:
                    # token-level says not foreign — discard primary result
                    continue

        if not lang_result or not is_foreign(lang_result, detect_text):
            continue

        # ADD THIS BLOCK
        # token_result = detect_language_by_tokens(detect_text, model_path)
        #
        # if token_result:
        #     if token_result["language"] != lang_result["language"]:
        #         continue  # disagreement → skip

        foreign_records.append({
            "source_type":          "paragraph",
            "word":                 text,
            "detect_word":          detect_text,
            "source_language":      lang_result["language"],
            "language_confidence":  lang_result["confidence"],
            "paragraph_no":         item.get("paragraph_no"),
            "x0":                   None,
            "y0":                   None,
            "x1":                   None,
            "y1":                   None,
        })

    # PDF + DOCX table cell-wise
    for item in parsed_data.get("table_data", []):
        text = item.get("cell_text", "").strip()
        detect_text = prepare_for_detection(text)
        if not detect_text:
            continue

        lang_result = detect_language(detect_text, model_path)
        if not lang_result:
            continue
        #print("lang result",lang_result)
        # Always run fallback as cross-check, not just on failure
        if not lang_result:
            continue
        # STEP 1 — trust strong primary result
        if is_foreign(lang_result, detect_text):
            pass  # keep it
        else:
            fallback_text = _extract_detectable_segment(detect_text)

            if fallback_text != detect_text:
                fallback_result = detect_language(fallback_text, model_path)
                if fallback_result and is_foreign(fallback_result, fallback_text):
                    lang_result = fallback_result
                    detect_text = fallback_text
                else:
                    token_result = detect_language_by_tokens(detect_text, model_path)
                    if token_result and is_foreign(token_result, detect_text):
                        # token-level agrees there's something foreign
                        lang_result = token_result
                    else:
                        # both fallbacks say not foreign — trust them over primary
                        continue
            else:
                token_result = detect_language_by_tokens(detect_text, model_path)
                if token_result and is_foreign(token_result, detect_text):
                    lang_result = token_result
                else:
                    # token-level says not foreign — discard primary result
                    continue

        if not lang_result or not is_foreign(lang_result, detect_text):
            continue
        # ADD THIS BLOCK
        # token_result = detect_language_by_tokens(detect_text, model_path)
        #
        # if token_result:
        #     if token_result["language"] != lang_result["language"]:
        #         continue  # disagreement → skip

        foreign_records.append({
            "source_type":          "table",
            "word":                 text,
            "detect_word": detect_text,
            "source_language":      lang_result["language"],
            "language_confidence":  lang_result["confidence"],
            "page_no":              item.get("page"),
            "table_no":             item.get("table_no"),
            "row_no":               item.get("row_no"),
            "column_no":            item.get("column_no"),
            "x0":                   item.get("x0"),
            "y0":                   item.get("y0"),
            "x1":                   item.get("x1"),
            "y1":                   item.get("y1"),
        })

    # TXT line-wise
    for line_no, text in enumerate(parsed_data.get("lines", []), start=1):
        text = text.strip()
        detect_text = prepare_for_detection(text)

        if not detect_text:
            continue

        lang_result = detect_language(detect_text, model_path)
        if not lang_result:
            continue
        # print("lang result",lang_result)
        # Always run fallback as cross-check, not just on failure
        # STEP 1 — trust strong primary result
        if is_foreign(lang_result, detect_text):
            pass  # keep it
        else:
            fallback_text = _extract_detectable_segment(detect_text)

            if fallback_text != detect_text:
                fallback_result = detect_language(fallback_text, model_path)
                if fallback_result and is_foreign(fallback_result, fallback_text):
                    lang_result = fallback_result
                    detect_text = fallback_text
                else:
                    token_result = detect_language_by_tokens(detect_text, model_path)
                    if token_result and is_foreign(token_result, detect_text):
                        # token-level agrees there's something foreign
                        lang_result = token_result
                    else:
                        # both fallbacks say not foreign — trust them over primary
                        continue
            else:
                token_result = detect_language_by_tokens(detect_text, model_path)
                if token_result and is_foreign(token_result, detect_text):
                    lang_result = token_result
                else:
                    # token-level says not foreign — discard primary result
                    continue

        if not lang_result or not is_foreign(lang_result, detect_text):
            continue
        #
        # # ADD THIS BLOCK
        # token_result = detect_language_by_tokens(detect_text, model_path)
        #
        # if token_result:
        #     if token_result["language"] != lang_result["language"]:
        #         continue  # disagreement → skip

        foreign_records.append({
            "source_type":          "txt_line",
            "word":                 text,
            "detect_word": detect_text,
            "source_language":      lang_result["language"],
            "language_confidence":  lang_result["confidence"],
            "line_no":              line_no,
            "x0":                   None,
            "y0":                   None,
            "x1":                   None,
            "y1":                   None,
        })

    return foreign_records

# ---------------------------------------------------------------------------
# BBox union helper  (FIX 6: was missing)
# ---------------------------------------------------------------------------

def union_bbox(bboxes) -> BboxPdf:
    bboxes = list(bboxes)
    if not bboxes:
        return BboxPdf(0, 0, 0, 0)
    x0 = min(b.x0 for b in bboxes)
    y0 = min(b.y0 for b in bboxes)
    x1 = max(b.x1 for b in bboxes)
    y1 = max(b.y1 for b in bboxes)
    return BboxPdf(x0, y0, x1, y1)


def detect_document_language(parsed_data: dict, model_path: str = r"D:\TMS\Models\lid.176.bin") -> str:
    all_text = []
    for item in parsed_data.get("text_data", []):
        t = item.get("text", "").strip()
        if len(t) > 30:
            all_text.append(t)
    for item in parsed_data.get("paragraph_data", []):
        t = item.get("text", "").strip()
        if len(t) > 30:
            all_text.append(t)
    for item in parsed_data.get("table_data", []):
        t = item.get("cell_text", "").strip()
        if len(t) > 30:
            all_text.append(t)
    if not all_text:
        return "en"
    sample = " ".join(all_text[:30])
    result = detect_language(sample, model_path)
    return result["language"] if result else "en"



def keep_primary_language_only(
    foreign_records: List[Dict],
) -> Tuple[List[Dict], Optional[str], Counter]:
    """
    Keep only phrases from the dominant (most frequent) language.
    Returns: (filtered_records, primary_lang, counts)
    """

    if not foreign_records:
        return [], None, Counter()

    counts = Counter(
        r.get("source_language")
        for r in foreign_records
        if r.get("source_language")
    )

    if not counts:
        return foreign_records, None, counts

    # 🔴 Only top 1 language
    primary_lang, _ = counts.most_common(1)[0]

    filtered = [
        r for r in foreign_records
        if r.get("source_language") == primary_lang
    ]

    return filtered, primary_lang, counts

english_dict = enchant.Dict("en_US")
translator = Translator()
def segment_foreign_record(record: dict) -> List[dict]:
    text = record.get("detect_word") or record.get("word", "")
    # print("detect   TEXTTTTTTT --========",text)
    original_text = record.get("word") or record.get("detected_word", "")
    # print("original   TEXTTTTTTT --========",original_text)
    highlight_text = normalize_for_detection(original_text)
    highlight_text = re.sub(r"\b\d+\s*mg/mL\b", " ", highlight_text, flags=re.I)
    highlight_text = re.sub(r"\s+", " ", highlight_text).strip()

    if not text:
        return []
    segments = []
    for token in highlight_text.split():
        clean = token.strip(".,;:!?\"'()[]{}/")
        if not clean or clean.lower() in IGNORE_TERMS or clean.lower() in COMMON_MEDICAL_TERMS:
            continue
        if _NUMBER_TOKEN_RE.match(clean):  # skip numbers and codes
            continue


        if len(clean) > 8:
            # This is an English word - skip it
            detection = translator.detect(clean)
            if detection.lang == "en":
                print("this might be english", clean)
                continue
        # sub_tokens = re.split(r"/", clean)
        # for sub in sub_tokens:
        #     sub = sub.strip(".,;:!?\"'()[]{}")  # re-strip after split
        #     if not sub or sub.lower() in IGNORE_TERMS or sub.lower() in COMMON_MEDICAL_TERMS:
        #         continue
        #     if _NUMBER_TOKEN_RE.match(sub):
        #         continue
        segments.append({
            "word":                clean,
            "original_token":      token,
            "source_type":         record.get("source_type"),
            "source_language":     record.get("source_language"),
            "language_confidence": record.get("language_confidence"),
            "page_no":             record.get("page_no"),
            "table_no":            record.get("table_no"),
            "row_no":              record.get("row_no"),
            "column_no":           record.get("column_no"),
            "paragraph_no":        record.get("paragraph_no"),
            "line_no":             record.get("line_no"),
            # bbox inherited from the cell/span — enriched to word-level later for PDFs
            "x0":                  record.get("x0"),
            "y0":                  record.get("y0"),
            "x1":                  record.get("x1"),
            "y1":                  record.get("y1"),
        })
    return segments


def segment_all_foreign_records(foreign_records: List[dict]) -> List[dict]:
    all_segments = []
    for record in foreign_records:
        all_segments.extend(segment_foreign_record(record))
    return all_segments


def enrich_segments_with_word_bboxes(segments: List[dict], file_path: str) -> List[dict]:
    """
    After language detection (which needs full cell text) and segmentation
    (which splits cells into word tokens), this function replaces the coarse
    cell/span bbox on each segment with a tight per-word bbox sourced from
    fitz get_text("words").

    Pipeline order:
        parse_pdf()                          <- cell-level text + cell bbox
        extract_foreign_content()            <- language detection on full cell text
        segment_all_foreign_records()        <- splits cell text into word tokens,
                                                each inheriting the cell bbox
        enrich_segments_with_word_bboxes()   <- THIS: replaces cell bbox → word bbox
    """
    if not file_path.lower().endswith(".pdf"):
        return segments

    doc = fitz.open(file_path)

    # Lazy cache: page_no (1-based) → list of fitz word tuples
    # fitz word tuple: (x0, y0, x1, y1, word, block_no, line_no, word_no)
    page_words: Dict[int, list] = {}

    def get_page_words(page_no: int) -> list:
        if page_no not in page_words:
            page_words[page_no] = doc[page_no - 1].get_text("words")
        return page_words[page_no]

    def strip_punct(s: str) -> str:
        return re.sub(r"^[\W]+|[\W]+$", "", s, flags=re.UNICODE)

    enriched = []
    last_word_index_by_page = {}
    for seg in segments:
        # Only attempt enrichment for PDF segments that have a bbox
        if (
            seg.get("source_type") not in ("table", "text")
            or seg.get("page_no") is None
            or seg.get("x0") is None
        ):
            enriched.append(seg)
            continue

        match_token = seg.get("original_token") or seg.get("word", "")

        token_norm = strip_punct(match_token).lower()
        page_no    = seg["page_no"]
        cell_x0, cell_y0 = seg["x0"], seg["y0"]
        cell_x1, cell_y1 = seg["x1"], seg["y1"]

        tol = 2.0  # points — accounts for coordinate rounding between pdfplumber/fitz
        best = None

        page_word_list = get_page_words(page_no)
        start_idx = last_word_index_by_page.get(page_no, 0)

        for idx in range(start_idx, len(page_word_list)):
            fw = page_word_list[idx]
            fx0, fy0, fx1, fy1, fword = fw[0], fw[1], fw[2], fw[3], fw[4]

            # optional bbox filter
            if fx0 < cell_x0 - tol or fy0 < cell_y0 - tol:
                continue
            if fx1 > cell_x1 + tol or fy1 > cell_y1 + tol:
                continue

            if strip_punct(fword).lower() == token_norm:
                best = fw
                last_word_index_by_page[page_no] = idx + 1
                break

        if best is not None:
            seg = {**seg, "x0": best[0], "y0": best[1], "x1": best[2], "y1": best[3]}

        if best is None:
            print(
                "MISS:",
                match_token,
                "| PAGE:", page_no,
                "| ROW:", seg.get("row_no"),
                "| CELL BBOX:", (cell_x0, cell_y0, cell_x1, cell_y1)
            )

        enriched.append(seg)


    doc.close()
    return enriched

def main():
    input_file  = r"D:\TMS\static\German-samples-Goethe-Institute.pdf"
    model_path  = r"D:\TMS\Models\lid.176.bin"

    parsed = process_file(input_file)
    document_lang = detect_document_language(parsed, model_path)
    print(f"Detected document language : {document_lang}")

    foreign_records = extract_foreign_content(parsed, base_language=document_lang)
    print(f"Foreign phrases found      : {len(foreign_records)}")

    foreign_records, top_lang, lang_counts = keep_primary_language_only(foreign_records)

    print("Top language:", top_lang)
    print("Counts:", dict(lang_counts))
    print("After filter:", len(foreign_records))

    segments = segment_all_foreign_records(foreign_records)
    # Enrich table/text segments with tight per-word bboxes.
    # Language detection already happened on full cell text, so this is safe.
    segments = enrich_segments_with_word_bboxes(segments, input_file)
    print(f"Word-level segments        : {len(segments)}")

    for seg in segments[:]:
        bbox_str = (
            f"bbox=({seg['x0']:.1f},{seg['y0']:.1f},{seg['x1']:.1f},{seg['y1']:.1f})"
            if seg.get("x0") is not None
            else "bbox=None"
        )
        print(
            f"  [{seg['source_type']:<10}] "
            f"{seg['word']:<25} "
            f"lang={seg['source_language']} "
            f"page={seg['page_no']} row={seg['row_no']} col={seg['column_no']} "
            f"{bbox_str}"
        )

    meta   = build_document_metadata(input_file, "pdf", document_lang)
    doc_id = meta["document_id"]
    print(f"\ndocument_id: {doc_id}")


if __name__ == "__main__":
    main()